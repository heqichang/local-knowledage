from abc import ABC, abstractmethod
from pathlib import Path


class BaseParser(ABC):
    @abstractmethod
    def parse(self, file_path: Path) -> str:
        pass


class TextParser(BaseParser):
    def parse(self, file_path: Path) -> str:
        try:
            return file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return file_path.read_text(encoding="gbk", errors="ignore")


class MarkdownParser(TextParser):
    pass


class PDFParser(BaseParser):
    def parse(self, file_path: Path) -> str:
        try:
            import fitz
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            return text
        except ImportError:
            return ""


class DocxParser(BaseParser):
    def parse(self, file_path: Path) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
            return text
        except ImportError:
            return ""


class ExcelParser(BaseParser):
    def parse(self, file_path: Path) -> str:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(file_path, read_only=True, data_only=True)
            text = ""
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text += f"=== 工作表: {sheet_name} ===\n"
                for row in sheet.iter_rows(values_only=True):
                    row_text = " ".join(str(cell) if cell is not None else "" for cell in row)
                    text += row_text + "\n"
            return text
        except ImportError:
            return ""


_parsers: dict[str, type[BaseParser]] = {
    "txt": TextParser,
    "md": MarkdownParser,
    "pdf": PDFParser,
    "docx": DocxParser,
    "doc": DocxParser,
    "xlsx": ExcelParser,
    "xls": ExcelParser,
}


def get_file_parser(file_type: str) -> BaseParser:
    parser_class = _parsers.get(file_type.lower(), TextParser)
    return parser_class()
